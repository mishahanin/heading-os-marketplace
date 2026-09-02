"""Shared utilities for rendering 31C corporate documents.

Consumed by: scripts/render-doctype.py and any skill that produces letters,
proposals, partnership documents, official documents, or xPagers.

The renderer resolves placeholders in locked HTML templates stored under
`datastore/brand/templates/doctypes/`, embeds brand assets (logos, fonts,
CSS), and writes output files. PDF rendering is delegated to
`scripts/html-to-pdf.py`. DOCX rendering uses python-docx.

Public API:
    TEMPLATE_REGISTRY - maps doctype to metadata
    render_html(doctype, data, workspace_root) -> str
    embed_brand_assets(html, workspace_root) -> str
    build_docx(doctype, data, out_path, workspace_root) -> Path
    validate_required_fields(doctype, data) -> list[str]
"""

from __future__ import annotations

import base64
import json
import re
import sys
from pathlib import Path


TEMPLATE_REGISTRY = {
    "letter": {
        "template": "letter.html",
        "formats": ["pdf", "docx"],
        "required": [
            "SENDER_NAME", "SENDER_TITLE", "SENDER_EMAIL", "SENDER_PHONE",
            "RECIPIENT_NAME", "RECIPIENT_TITLE", "RECIPIENT_ORG",
            "SUBJECT", "DATE", "REF_ID", "SALUTATION", "BODY_HTML",
        ],
    },
    "proposal": {
        "template": "proposal.html",
        "formats": ["pdf", "docx"],
        "required": [
            "SENDER_NAME", "SENDER_TITLE", "SENDER_EMAIL", "SENDER_PHONE",
            "RECIPIENT_NAME", "RECIPIENT_TITLE", "RECIPIENT_ORG", "RECIPIENT_COUNTRY",
            "SUBJECT", "LEDE", "DATE", "REF_ID",
            "EXECUTIVE_OPENING_HTML", "OPPORTUNITY_HTML", "SOLUTION_HTML",
            "PROOF_HTML", "COMMERCIAL_INTRO_HTML", "NEXT_STEPS_HTML",
            "PRICING_LINES",
        ],
    },
    "partnership": {
        "template": "partnership.html",
        "formats": ["pdf", "docx"],
        "required": [
            "SUBTYPE", "SUBJECT", "REF_ID", "EFFECTIVE_DATE", "TERM",
            "PARTY_A_NAME", "PARTY_A_SHORT", "PARTY_A_ENTITY_TYPE",
            "PARTY_A_JURISDICTION", "PARTY_A_ADDRESS",
            "PARTY_A_SIGNATORY_NAME", "PARTY_A_SIGNATORY_TITLE",
            "PARTY_B_NAME", "PARTY_B_SHORT", "PARTY_B_ENTITY_TYPE",
            "PARTY_B_JURISDICTION", "PARTY_B_ADDRESS",
            "PARTY_B_SIGNATORY_NAME", "PARTY_B_SIGNATORY_TITLE",
            "PURPOSE_HTML", "SCOPE_HTML", "GOVERNANCE_HTML", "CLAUSES",
        ],
    },
    "official": {
        "template": "official.html",
        "formats": ["pdf", "docx"],
        "required": [
            "CLASS", "REF_ID", "DATE", "PLACE",
            "ISSUER_NAME", "ISSUER_TITLE", "SUBJECT",
        ],
    },
    "xpager": {
        "template": "xpager.html",
        "formats": ["pdf", "html"],
        "required": [
            "PRODUCT_NAME", "PRODUCT_BASE", "PRODUCT_SUFFIX", "TAGLINE", "DATE",
            "STATS", "CAPABILITY_HEADLINE", "CAPABILITY_SUB", "CAPABILITIES",
            "PROOF_HEADLINE", "PROOF_SUB", "PROOF_POINTS",
            "CONTACT_NAME", "CONTACT_TITLE", "CONTACT_EMAIL", "CONTACT_PHONE",
        ],
    },
}


def _under_datastore(rel_path: str) -> Path | None:
    """Map a `datastore/...` relative path onto the resolved datastore root.

    None when `rel_path` is not under `datastore/`, or when the seam cannot be
    resolved at all - a public clone with no overlay, for instance. The failure
    is printed rather than swallowed, because a silent None here reappears
    downstream as a bare "template not found" that names the wrong cause.
    """
    prefix = "datastore/"
    if not rel_path.startswith(prefix):
        return None
    try:
        from scripts.utils.workspace import get_datastore_dir
        return get_datastore_dir() / rel_path[len(prefix):]
    except Exception as exc:  # noqa: BLE001 - reported, never hidden
        print(f"doctype_renderer: cannot resolve the datastore root "
              f"({type(exc).__name__}: {exc}); falling back to the workspace tree.",
              file=sys.stderr)
        return None


def _resolve_under_corporate(workspace_root: Path, rel_path: str) -> Path:
    """Resolve a path that lives at workspace root on CEO master and under
    corporate/ on exec workspaces (where shared content is sync-mirrored).

    The third candidate is the data-root seam, and it is the one that works in
    production. `scripts/render-doctype.py` passes `get_workspace_root()`, the
    ENGINE clone, while `datastore/brand/` lives in the private DATA overlay - so
    every real call resolved to `<engine>/datastore/brand/...`, which does not
    exist, and all five locked doctypes (`/corporate-letter`, `/proposal`,
    `/partnership-doc`, `/official-doc`, `/xpager`) died with FileNotFoundError.

    The explicit-root candidates stay FIRST so a caller that hands over a real
    tree still gets that tree, which is what the fixtures rely on.
    """
    direct = workspace_root / rel_path
    if direct.exists():
        return direct
    under_corporate = workspace_root / "corporate" / rel_path
    if under_corporate.exists():
        return under_corporate
    seam = _under_datastore(rel_path)
    if seam is not None and seam.exists():
        return seam
    return direct


def _templates_dir(workspace_root: Path) -> Path:
    return _resolve_under_corporate(workspace_root, "datastore/brand/templates/doctypes")


def _assets_dir(workspace_root: Path) -> Path:
    return _resolve_under_corporate(workspace_root, "datastore/brand/assets")


def _fonts_dir(workspace_root: Path) -> Path:
    return _resolve_under_corporate(workspace_root, "datastore/brand/fonts/GT Standard")


def _inter_dir(workspace_root: Path) -> Path:
    """The Cyrillic fallback family, a sibling of `GT Standard`, not a child.

    `_resolve_brand_assets` used to build this as `_fonts_dir(root) / "Inter"`,
    and `_fonts_dir` already ends in `GT Standard`. The resulting
    `.../fonts/GT Standard/Inter` has never existed, so both Inter faces embedded
    as the empty string and every Russian run fell through to a system font.
    Resolved through the same seam as its sibling rather than as
    `_fonts_dir(root).parent`, so the two lookups stay independent.
    """
    return _resolve_under_corporate(workspace_root, "datastore/brand/fonts/Inter")


def validate_required_fields(doctype: str, data: dict) -> list[str]:
    """Return a list of missing required fields."""
    if doctype not in TEMPLATE_REGISTRY:
        return [f"unknown_doctype:{doctype}"]
    required = TEMPLATE_REGISTRY[doctype]["required"]
    missing = []
    for field in required:
        value = data.get(field)
        if value is None or isinstance(value, str) and not value.strip() or isinstance(value, list) and not value:
            missing.append(field)
    return missing


def _encode_file_b64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode("ascii")


def _docx_text(html) -> str:
    """The HTML body of a document, as the plain text its DOCX twin carries.

    `build_docx` used to hold a four-line private stripper that turned `<br>`
    and `</p><p>` into newlines and then deleted every other tag with NO
    separator: `<ul><li>A</li><li>B</li></ul>` arrived as `AB`, a table row as
    `ab`, and `&amp;` as four literal characters. The PDF was correct, because
    the template gets the HTML raw, so nothing the operator saw showed what the
    counterparty would open in the editable copy.

    `html_text.strip_html` is the shared parser, was rewritten for this exact
    fusion, and its docstring asks a new caller to import rather than copy it.

    `sanitize` runs afterwards and is NOT optional: the shared parser decodes
    entities, `&nbsp;` decodes to U+00A0, and `.claude/rules/hidden-chars.md`
    bans that character in every generated artifact. Adopting the parser without
    this line would trade a fused word for a banned one.
    """
    from scripts.utils.html_text import strip_html
    from scripts.utils.sanitize_text import sanitize

    return sanitize(strip_html(html))


def _docx_paragraphs(html) -> list[str]:
    """The same text, split where the block elements were.

    One paragraph per line, empties dropped, so a list becomes one DOCX
    paragraph per item instead of one paragraph per `</p><p>` boundary. The old
    `.split("\\n\\n")` could only ever see that one boundary, which is why every
    other block element fused.
    """
    return [line.strip() for line in _docx_text(html).split("\n") if line.strip()]


def _embed_asset(path: Path, mime: str) -> str:
    """Return a data: URI for the asset, or "" when the file is not there.

    The miss is REPORTED. Returning "" and saying nothing is how a wrong Inter
    path survived: the template got `src: url("")`, Chromium fell back to a
    system font, and the render still exited 0 with `[WROTE]` printed. By the
    time this runs, `render_html` has already read the locked template out of
    the same brand tree, so a file missing here is an anomaly rather than a
    public clone with no overlay. "" is still returned, because a half-branded
    document beats no document at all when the operator is mid-deadline.
    """
    if not path.exists():
        print(f"doctype_renderer: brand asset not found, rendering without it: {path}",
              file=sys.stderr)
        return ""
    b64 = _encode_file_b64(path)
    return f"data:{mime};base64,{b64}"


def _resolve_brand_assets(workspace_root: Path) -> dict:
    """Embed every brand asset the locked templates reference.

    The filenames come from the private manifest, never from this file. The
    engine repository is public, and on 2026-09-02 the operator ruled that a
    datastore filename is itself private; `scripts/utils/brand_assets.py` carries
    the whole reasoning. Only the DIRECTORY lookups stay here, because
    `_resolve_under_corporate` has to keep trying an explicit `workspace_root`
    first for the fixtures that hand it a synthetic tree.
    """
    from scripts.utils.brand_assets import brand_asset_name, load_manifest

    manifest = load_manifest()
    logos = _assets_dir(workspace_root) / "logos"
    fonts = _fonts_dir(workspace_root)
    inter = _inter_dir(workspace_root)

    def logo(key: str) -> str:
        return _embed_asset(logos / brand_asset_name(key, manifest), "image/png")

    def face(directory: Path, key: str) -> str:
        return _embed_asset(directory / brand_asset_name(key, manifest), "font/ttf")

    return {
        "LOGO_BLUE": logo("logo_primary"),
        "LOGO_WHITE": logo("logo_on_dark"),
        "LOGO_BLACK": logo("logo_on_light"),
        # TTF (not WOFF2) so Chromium embeds the fonts in PDF as CIDFont Type 2
        # (TrueType subset with hinting) rather than as Type 3 outlines. Type 3
        # rendering thickens small solid glyphs — periods, commas, the middle
        # dot — making them read as bold against Light-weight letterforms.
        "FONT_LIGHT": face(fonts, "font_gt_m_light_ttf"),
        "FONT_MEDIUM": face(fonts, "font_gt_m_medium_ttf"),
        # Inter Light + Medium (SIL OFL) — Cyrillic fallback. GT Standard has
        # no Cyrillic glyphs; without these, Russian text falls back to system
        # Segoe UI / Arial at a heavier weight than the Latin column. Static
        # TTFs (not the variable font) so Chromium embeds them as TrueType
        # subsets in PDF with hinting preserved; the variable font got
        # converted to Type 3 outlines for Cyrillic runs and the bold-weight
        # interpolation failed to a few characters fell through to Arial.
        "FONT_INTER_LIGHT": face(inter, "font_inter_light_ttf"),
        "FONT_INTER_MEDIUM": face(inter, "font_inter_medium_ttf"),
    }


def _load_brand_css(workspace_root: Path) -> str:
    css_path = _templates_dir(workspace_root) / "_shared" / "base.css"
    return css_path.read_text(encoding="utf-8")


# The two grammars have to agree. A scalar `{{PHASE_2}}` was always legal, but
# the section name was `[A-Z_]+` - no digits - so `{{#PHASE_2_ITEMS}}` matched
# nothing, `_render_sections` left the block alone, and `_substitute_scalars`
# could not touch it either because `{{#...}}` and `{{/...}}` carry a `#` and a
# `/`. The raw template syntax and the whole loop body were then emitted
# verbatim into the rendered HTML and the PDF, with a successful exit code.
# Measured 2026-08-30: rendering `{{#PHASE_2_ITEMS}}<li>{{.}}</li>{{/PHASE_2_ITEMS}}`
# returned that string unchanged.
_SECTION_RE = re.compile(r"\{\{#([A-Z_][A-Z0-9_]*)\}\}(.*?)\{\{/\1\}\}", re.DOTALL)
_VAR_RE = re.compile(r"\{\{([A-Z_][A-Z0-9_]*|\.)\}\}")


def _render_sections(template: str, data: dict) -> str:
    """Expand {{#LIST}}...{{/LIST}} blocks.

    - If the value is a list of dicts, repeat the inner block for each dict,
      substituting {{key}} with the dict's fields.
    - If the value is a list of strings, repeat with {{.}} replaced by each
      string.
    - If the value is a non-empty string/dict, render the inner block once with
      outer-scope substitution.
    - If the value is falsy/empty, drop the block.
    """
    def replace(match: re.Match) -> str:
        key = match.group(1)
        inner = match.group(2)
        value = data.get(key)
        if not value:
            return ""
        if isinstance(value, list):
            parts = []
            for item in value:
                if isinstance(item, dict):
                    piece = inner
                    for k, v in item.items():
                        piece = piece.replace("{{" + k + "}}", str(v))
                    parts.append(piece)
                else:
                    parts.append(inner.replace("{{.}}", str(item)))
            return "".join(parts)
        if isinstance(value, str):
            return inner.replace("{{.}}", value)
        return inner

    # Apply until no more section tags remain (handles nested sections).
    prev = None
    while prev != template:
        prev = template
        template = _SECTION_RE.sub(replace, template)
    return template


def _substitute_scalars(template: str, data: dict) -> str:
    def replace(match: re.Match) -> str:
        key = match.group(1)
        if key == ".":
            return match.group(0)
        if key in data:
            value = data[key]
            if isinstance(value, (str, int, float)):
                return str(value)
        return ""
    return _VAR_RE.sub(replace, template)


def render_html(doctype: str, data: dict, workspace_root: Path) -> str:
    """Render the HTML for a doctype, embedding brand assets and CSS."""
    if doctype not in TEMPLATE_REGISTRY:
        raise ValueError(f"Unknown doctype: {doctype}")
    template_path = _templates_dir(workspace_root) / TEMPLATE_REGISTRY[doctype]["template"]
    template = template_path.read_text(encoding="utf-8")

    brand_css = _load_brand_css(workspace_root)
    brand_assets = _resolve_brand_assets(workspace_root)

    # Substitute brand CSS into the CSS block first.
    template = template.replace("{{BRAND_CSS}}", brand_css)
    # Then substitute font placeholders inside the CSS.
    template = template.replace("{{FONT_LIGHT}}", brand_assets["FONT_LIGHT"])
    template = template.replace("{{FONT_MEDIUM}}", brand_assets["FONT_MEDIUM"])
    template = template.replace("{{FONT_INTER_LIGHT}}", brand_assets["FONT_INTER_LIGHT"])
    template = template.replace("{{FONT_INTER_MEDIUM}}", brand_assets["FONT_INTER_MEDIUM"])
    # Then logo placeholders wherever they appear.
    template = template.replace("{{LOGO_BLUE}}", brand_assets["LOGO_BLUE"])
    template = template.replace("{{LOGO_WHITE}}", brand_assets["LOGO_WHITE"])
    template = template.replace("{{LOGO_BLACK}}", brand_assets["LOGO_BLACK"])

    # Expand list/section blocks next (they may contain scalar vars).
    template = _render_sections(template, data)

    # Finally substitute all remaining scalar placeholders.
    template = _substitute_scalars(template, data)

    return template


def build_docx(doctype: str, data: dict, out_path: Path, workspace_root: Path) -> Path:
    """Build a DOCX for types that support it. Letters/proposals/partnerships/officials.

    Uses python-docx to create an editable Word document with the locked
    letterhead, body, signature, and footer.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed. Run: pip install python-docx"
        ) from exc

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "GT Standard"
    style.font.size = Pt(11)

    def add_heading(text: str, size: int = 18, bold: bool = True, color: tuple[int, int, int] | None = None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_label_row(label: str, value: str):
        p = doc.add_paragraph()
        lbl = p.add_run(f"{label.upper()}   ")
        lbl.font.size = Pt(8)
        lbl.font.bold = True
        lbl.font.color.rgb = RGBColor(0x5A, 0x5A, 0x78)
        val = p.add_run(value)
        val.font.size = Pt(10)

    # Header block
    add_heading("31 Concept", size=22, color=(0x15, 0x15, 0x15))
    tag = doc.add_paragraph()
    tag_run = tag.add_run("DEEP PACKET INTELLIGENCE")
    tag_run.font.size = Pt(9)
    tag_run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x78)
    doc.add_paragraph()  # spacer

    if doctype == "letter":
        add_label_row("From", f"{data['SENDER_NAME']}, {data['SENDER_TITLE']} - {data['SENDER_EMAIL']} - {data['SENDER_PHONE']}")
        add_label_row("To", f"{data['RECIPIENT_NAME']}, {data['RECIPIENT_TITLE']} - {data['RECIPIENT_ORG']}")
        add_label_row("Date", data["DATE"])
        add_label_row("Ref", data["REF_ID"])
        doc.add_paragraph()
        add_heading(data["SUBJECT"], size=14, color=(0x15, 0x15, 0x15))
        doc.add_paragraph()
        # Through the same helper as the body. This line used to strip `<p>` by
        # name, which is the private stripper's defect in miniature: any other
        # tag, and every entity, reached the page verbatim.
        doc.add_paragraph(_docx_text(data.get("SALUTATION", "")))
        for para in _docx_paragraphs(data["BODY_HTML"]):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        name_p = doc.add_paragraph()
        name_p.add_run(data["SENDER_NAME"]).font.bold = True
        doc.add_paragraph(f"{data['SENDER_TITLE']} - 31 Concept")
        doc.add_paragraph(f"{data['SENDER_EMAIL']} - {data['SENDER_PHONE']}")

    elif doctype == "proposal":
        add_label_row("Proposal", data["SUBJECT"])
        add_label_row("Prepared for", f"{data['RECIPIENT_ORG']} - {data['RECIPIENT_COUNTRY']}")
        add_label_row("Attention", f"{data['RECIPIENT_NAME']}, {data['RECIPIENT_TITLE']}")
        add_label_row("Date", data["DATE"])
        add_label_row("Ref", data["REF_ID"])
        doc.add_paragraph()
        add_heading("Executive Opening", size=14, color=(0x5B, 0x5F, 0xFF))
        for para in _docx_paragraphs(data["EXECUTIVE_OPENING_HTML"]):
            if para.strip():
                doc.add_paragraph(para.strip())
        add_heading("The Opportunity", size=14, color=(0x5B, 0x5F, 0xFF))
        for para in _docx_paragraphs(data["OPPORTUNITY_HTML"]):
            if para.strip():
                doc.add_paragraph(para.strip())
        add_heading("Solution Structure", size=14, color=(0x5B, 0x5F, 0xFF))
        for para in _docx_paragraphs(data["SOLUTION_HTML"]):
            if para.strip():
                doc.add_paragraph(para.strip())
        add_heading("Why 31C", size=14, color=(0x5B, 0x5F, 0xFF))
        for para in _docx_paragraphs(data["PROOF_HTML"]):
            if para.strip():
                doc.add_paragraph(para.strip())
        add_heading("Commercial Terms", size=14, color=(0x5B, 0x5F, 0xFF))
        for para in _docx_paragraphs(data["COMMERCIAL_INTRO_HTML"]):
            if para.strip():
                doc.add_paragraph(para.strip())
        for line in data.get("PRICING_LINES", []):
            p = doc.add_paragraph()
            p.add_run(f"{line['label']}: ").font.bold = True
            p.add_run(line["value"])
        add_heading("Next Steps", size=14, color=(0x5B, 0x5F, 0xFF))
        for para in _docx_paragraphs(data["NEXT_STEPS_HTML"]):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        name_p = doc.add_paragraph()
        name_p.add_run(data["SENDER_NAME"]).font.bold = True
        doc.add_paragraph(f"{data['SENDER_TITLE']} - 31 Concept")
        doc.add_paragraph(f"{data['SENDER_EMAIL']} - {data['SENDER_PHONE']}")

    elif doctype == "partnership":
        add_heading(data["SUBTYPE"], size=10, color=(0x5B, 0x5F, 0xFF))
        add_heading(data["SUBJECT"], size=16, color=(0x15, 0x15, 0x15))
        add_label_row("Ref", data["REF_ID"])
        add_label_row("Effective", data["EFFECTIVE_DATE"])
        add_label_row("Term", data["TERM"])
        doc.add_paragraph()
        add_heading("Between the Parties", size=12, color=(0x5B, 0x5F, 0xFF))
        doc.add_paragraph(
            f"{data['PARTY_A_NAME']}, a {data['PARTY_A_ENTITY_TYPE']} organised under the laws of "
            f"{data['PARTY_A_JURISDICTION']}, with its principal office at {data['PARTY_A_ADDRESS']} "
            f"(\"{data['PARTY_A_SHORT']}\")."
        )
        doc.add_paragraph(
            f"{data['PARTY_B_NAME']}, a {data['PARTY_B_ENTITY_TYPE']} organised under the laws of "
            f"{data['PARTY_B_JURISDICTION']}, with its principal office at {data['PARTY_B_ADDRESS']} "
            f"(\"{data['PARTY_B_SHORT']}\")."
        )
        add_heading("Purpose", size=12, color=(0x5B, 0x5F, 0xFF))
        doc.add_paragraph(_docx_text(data["PURPOSE_HTML"]))
        add_heading("Scope of Collaboration", size=12, color=(0x5B, 0x5F, 0xFF))
        doc.add_paragraph(_docx_text(data["SCOPE_HTML"]))
        for clause in data.get("CLAUSES", []):
            add_heading(f"Clause {clause['num']} - {clause['title']}", size=11, color=(0x5B, 0x5F, 0xFF))
            doc.add_paragraph(_docx_text(clause["body"]))
        add_heading("Governance and Dispute Resolution", size=12, color=(0x5B, 0x5F, 0xFF))
        doc.add_paragraph(_docx_text(data["GOVERNANCE_HTML"]))
        add_heading("Confidentiality", size=12, color=(0x5B, 0x5F, 0xFF))
        doc.add_paragraph(
            "Each party shall treat all non-public information exchanged under this "
            f"{data['SUBTYPE']} as confidential and shall not disclose it to any third party without "
            "the prior written consent of the disclosing party, except as required by law or regulation."
        )
        doc.add_paragraph()
        doc.add_paragraph("Executed on the Effective Date shown above.")
        doc.add_paragraph()
        sig_p = doc.add_paragraph()
        sig_p.add_run(f"For {data['PARTY_A_SHORT']}:").font.bold = True
        doc.add_paragraph(f"{data['PARTY_A_SIGNATORY_NAME']} - {data['PARTY_A_SIGNATORY_TITLE']}")
        doc.add_paragraph("Signature: ______________________    Date: ______________")
        doc.add_paragraph()
        sig_p = doc.add_paragraph()
        sig_p.add_run(f"For {data['PARTY_B_SHORT']}:").font.bold = True
        doc.add_paragraph(f"{data['PARTY_B_SIGNATORY_NAME']} - {data['PARTY_B_SIGNATORY_TITLE']}")
        doc.add_paragraph("Signature: ______________________    Date: ______________")

    elif doctype == "official":
        add_heading(data["CLASS"], size=10, color=(0x5B, 0x5F, 0xFF))
        add_heading(data["SUBJECT"], size=16, color=(0x15, 0x15, 0x15))
        add_label_row("Reference", data["REF_ID"])
        add_label_row("Date", data["DATE"])
        add_label_row("Place of Issue", data["PLACE"])
        add_label_row("Issued By", f"{data['ISSUER_NAME']}, {data['ISSUER_TITLE']}")
        if data.get("SUBJECT_LINE"):
            add_label_row("Subject", data["SUBJECT_LINE"])
        doc.add_paragraph()
        if data.get("PREAMBLE"):
            doc.add_paragraph(_docx_text(data["PREAMBLE"]))
        for w in data.get("WHEREAS_CLAUSES", []):
            p = doc.add_paragraph()
            p.add_run("Whereas ").font.bold = True
            p.add_run(_docx_text(w))
        for r in data.get("RESOLVED_BLOCKS", []):
            p = doc.add_paragraph()
            p.add_run("NOW, THEREFORE, RESOLVED: ").font.bold = True
            p.add_run(_docx_text(r))
        if data.get("CLOSING_HTML"):
            doc.add_paragraph(_docx_text(data["CLOSING_HTML"]))
        doc.add_paragraph()
        doc.add_paragraph("Executed on the date first written above.")
        doc.add_paragraph()
        sig_p = doc.add_paragraph()
        sig_p.add_run(data["ISSUER_NAME"]).font.bold = True
        doc.add_paragraph(f"{data['ISSUER_TITLE']} - 31 Concept")
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("OFFICIAL SEAL OF 31 CONCEPT").font.color.rgb = RGBColor(0x5A, 0x5A, 0x78)
    else:
        raise ValueError(f"DOCX not supported for doctype: {doctype}")

    # Footer confidentiality
    doc.add_paragraph()
    conf = doc.add_paragraph()
    conf_run = conf.add_run(
        "This document contains confidential and privileged information intended solely for the "
        "named recipient. Unauthorised review, disclosure, or distribution is prohibited."
    )
    conf_run.font.size = Pt(7)
    conf_run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x78)
    footer = doc.add_paragraph()
    footer_run = footer.add_run("31C.io - (c) 2025-2026 / 31 Concept / Proprietary & Confidential")
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x78)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def load_data(data_path: Path) -> dict:
    """Load placeholder data from a JSON file."""
    return json.loads(data_path.read_text(encoding="utf-8"))
